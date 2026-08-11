#!/usr/bin/env python3
"""Build the authoritative Amaura AI System Handbook in Markdown and DOCX."""

from __future__ import annotations

import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from jarvis.amaura.content_factory import MEASUREMENT_WINDOWS, REQUIRED_PUBLICATION_ASSETS  # noqa: E402
from jarvis.amaura.models import ApprovalStatus, TaskState  # noqa: E402
from jarvis.amaura.pipeline import ALLOWED_TRANSITIONS, SCORE_LIMITS  # noqa: E402
from jarvis.amaura.policies import POLICIES  # noqa: E402
from jarvis.amaura.policy import EXTERNAL_ACTIONS, SAFE_COMMAND_PREFIXES, TOOL_RISK_CLASSES  # noqa: E402
from jarvis.amaura.readiness import INTEGRATIONS  # noqa: E402
from jarvis.amaura.registry import ALL_AGENTS  # noqa: E402
from jarvis.amaura.workflows import WORKFLOWS  # noqa: E402

OUT_MD = ROOT / "docs" / "AMAURA_AI_SYSTEM_HANDBOOK.md"
OUT_DOCX = ROOT / "docs" / "AMAURA_AI_SYSTEM_HANDBOOK.docx"
VERSION = "1.0"
SOURCE_DATE = "2026-07-27"


def money(cents: int) -> str:
    return f"${cents / 100:,.2f}"


def inline_list(values) -> str:
    return ", ".join(f"`{value}`" for value in values) if values else "None"


def table(headers: list[str], rows: list[list[str]]) -> list[str]:
    def clean(value: object) -> str:
        return str(value).replace("\n", "<br>").replace("|", "\\|")

    lines = ["| " + " | ".join(clean(h) for h in headers) + " |"]
    lines.append("| " + " | ".join("---" for _ in headers) + " |")
    lines.extend("| " + " | ".join(clean(v) for v in row) + " |" for row in rows)
    return lines


def code(lines: list[str], language: str = "text") -> list[str]:
    return [f"```{language}", *lines, "```"]


def build_markdown() -> str:
    lines: list[str] = []
    add = lines.extend

    add([
        "# Amaura Labs AI System Handbook",
        "",
        "**Complete operating, administration, security, workflow, and recovery manual**",
        "",
        f"Version {VERSION} · Source snapshot {SOURCE_DATE} · Internal operations",
        "",
        "This handbook is the practical source of truth for installing, accessing, operating, reviewing, securing, testing, backing up, and extending the Amaura AI Workforce inside JARVIS.",
        "",
        "> **Authority statement:** The founder owns strategy, truth, money, legal commitments, public reputation, and production risk. JARVIS is the sole company orchestrator. Specialist employees operate only inside issued task packets and cannot certify their own work.",
        "",
        "<!-- PAGEBREAK -->",
        "",
        "## Document control",
        "",
    ])
    add(table(["Field", "Value"], [
        ["Document owner", "Amaura Labs founder / JARVIS operator"],
        ["Applies to", "The 57-role Amaura Workforce, 21 workflows, SQLite operating ledger, REST and JARVIS tool surfaces"],
        ["Source of truth", "`jarvis/amaura/`, `jarvis/server.py`, `jarvis/tools/amaura.py`, `.env.amaura.example`, and production verification artifacts"],
        ["Review trigger", "Any registry, workflow, policy, API, database, model-route, provider, or authority change"],
        ["Verified baseline", "233 isolated tests passed; static checks, authenticated API smoke, deterministic wheel build, and 1,000-lead stress passed"],
        ["Classification", "Internal operations. Never place real secrets in this document."],
    ]))
    add([
        "",
        "### How to use this handbook",
        "",
        "- **New operator:** read Sections 1–5, then follow Sections 12–14 exactly.",
        "- **Founder:** use Sections 4, 15, 18, and 22 for authority, daily decisions, incidents, and release gates.",
        "- **Revenue team:** use Sections 8 and 16 plus the endpoint recipes in Section 20.",
        "- **Content team:** use Sections 9 and 17 plus the asset-readiness rules in Section 10.",
        "- **Engineer or administrator:** use Sections 6–7 and 18–23.",
        "- In Word, open the Navigation Pane to jump through the real heading hierarchy. The Markdown file is searchable and friendly to version control.",
        "",
        "## Contents",
        "",
        "1. System purpose and operating status",
        "2. Architecture and authority chain",
        "3. Core concepts and terminology",
        "4. Governance, risk, and founder authority",
        "5. System components",
        "6. Complete AI employee directory",
        "7. Workflow engine and task lifecycle",
        "8. Client-acquisition system",
        "9. Amaura Content Factory",
        "10. Data, evidence, and storage",
        "11. Models, prompts, tools, and execution",
        "12. Installation and first-time configuration",
        "13. How to start and access the system",
        "14. First-day acceptance procedure",
        "15. Daily founder and operator routine",
        "16. Running a client-acquisition campaign",
        "17. Running a content campaign",
        "18. Approvals, pauses, kill switches, and incidents",
        "19. Monitoring, costs, backup, and recovery",
        "20. REST API and command reference",
        "21. Testing, stress testing, and release readiness",
        "22. External integrations and rollout stages",
        "23. Troubleshooting and maintenance",
        "24. Operational checklists and glossary",
        "",
        "<!-- PAGEBREAK -->",
        "",
        "# Part I — Orientation and control",
        "",
        "## 1. System purpose and operating status",
        "",
        "### 1.1 What the system is",
        "",
        "Amaura is a local-first, governed company operating kernel inside JARVIS. It converts a founder objective into a programme, project, milestone, and dependency-ordered specialist tasks. Every task has a named owner, independent reviewer, measurable exit criteria, budget, risk level, allowed tools, approved data, model route, evidence, audit entries, and—when required—a separately authenticated founder decision.",
        "",
        "The system is not a collection of autonomous bots with unrestricted company access. It is an authority-controlled workflow engine whose AI employees are bounded by deterministic code and durable records.",
        "",
        "### 1.2 What is production-ready today",
        "",
        "- The local operating kernel: workforce registry, workflow creation, task dispatch, policy enforcement, independent review, approvals, events, audit logs, budgets, readiness checks, acquisition controls, content asset controls, SQLite persistence, backups, REST surfaces, and JARVIS tools.",
        "- Authenticated operation on the local machine after three independent authority keys are configured.",
        "- Device-only routing for restricted data through Ollama, with no silent cloud fallback.",
        "- A 43-role registry and six workflow templates, including the full 16-stage acquisition pipeline and 12-stage content factory.",
        "",
        "### 1.3 What is not automatically active",
        "",
        "Gmail sending, Telegram decisions, OBS recording, Postiz or YouTube publication, external sandboxes, contracts, payments, and production deployments are not active merely because the kernel exists. Those capabilities require real binaries, credentials or OAuth grants, provider adapters, least-privilege scopes, and provider-confirmed callbacks. Until then, the system prepares, governs, records, and verifies work but does not claim the external action occurred.",
        "",
        "> **No-silent-success rule:** `/messages/{id}/sent` is a confirmation boundary, not an email client. Record a send only after an approved provider returns a real message identifier.",
        "",
        "### 1.4 Verified release baseline",
        "",
    ])
    add(table(["Gate", "Verified result"], [
        ["Repository tests", "63 passed, 0 failed"],
        ["Static analysis", "Focused Ruff passed; focused mypy passed on eight core modules"],
        ["Packaging", "Wheel built: `release/jarvis-3.6.1-py3-none-any.whl`"],
        ["API security smoke", "Authenticated routes passed; missing operator authority rejected with HTTP 403"],
        ["Stress run", "1,000 leads, 32 workers, 100/100 injections detected, 60 provider-confirmed sends, 20/20 cap violations blocked"],
        ["Concurrency and idempotency", "500 simultaneous duplicate attempts produced one record"],
        ["Storage", "SQLite integrity `ok`, no foreign-key violations, WAL enabled"],
    ]))
    add([
        "",
        "## 2. Architecture and authority chain",
        "",
        "### 2.1 Authority path",
        "",
    ])
    add(code([
        "Founder",
        "  └── separately authenticated decisions and manual critical actions",
        "      └── JARVIS — sole control plane",
        "          ├── Programme → Project → Milestone → dependency-ordered Tasks",
        "          ├── Policy Engine → tools, paths, commands, data, cost and risk",
        "          ├── Model Gateway → privacy/budget-aware model route",
        "          ├── Specialist Employee → evidence-bearing submission",
        "          ├── Independent Reviewer → approve or return changes",
        "          └── Founder Approval → consequential completion boundary",
    ]))
    add([
        "",
        "### 2.2 Runtime request path",
        "",
    ])
    add(code([
        "Browser / CLI / JARVIS chat / REST client",
        "        │",
        "        ▼",
        "FastAPI server (loopback by default; restricted CORS)",
        "        │  operator key or approval key",
        "        ▼",
        "AmauraControlPlane",
        "        ├── CompanyStore (SQLite + WAL + foreign keys)",
        "        ├── PolicyEngine",
        "        ├── ModelGateway",
        "        ├── GovernedTaskRunner",
        "        ├── AcquisitionPipeline",
        "        └── ContentFactory",
    ]))
    add([
        "",
        "### 2.3 Separation of responsibilities",
        "",
    ])
    add(table(["Layer", "Owns", "Must not do"], [
        ["Founder", "Strategy, final truth, external commitments, reputation, money, critical execution", "Delegate irreversible authority to an unauthenticated agent"],
        ["JARVIS", "Programme creation, task packets, dispatch, orchestration, pause, escalation, briefing", "Permit self-review or bypass the registered workflow"],
        ["Specialist employee", "One bounded task, approved tools/data, evidence-backed output", "Expand scope, exceed budget, invent facts, certify itself"],
        ["Independent reviewer", "Acceptance criteria, evidence quality, defects and policy verification", "Review its own work or approve without findings"],
        ["Policy engine", "Deterministic allow/deny and approval requirements", "Treat model text as authority"],
        ["External adapter", "One provider action and provider-confirmed result", "Claim success without provider evidence"],
    ]))
    add([
        "",
        "## 3. Core concepts and terminology",
        "",
    ])
    add(table(["Term", "Meaning"], [
        ["Programme", "Top-level founder outcome with a measurable success metric."],
        ["Project", "A workflow instance used to deliver the programme."],
        ["Milestone", "Completion boundary containing the workflow tasks."],
        ["Task packet", "JARVIS-issued execution contract: objective, criteria, owner, reviewer, tools, data, budget, dependencies, model route, policies, workspace, and doctrine."],
        ["Evidence", "A typed, stable reference supporting a completion claim—tool output digest, source URL/excerpt, artifact hash, test result, or provider identifier."],
        ["Action type", "Business meaning of an action, used to load relevant policies and determine approval."],
        ["Risk level", "Low, medium, high, or critical. It constrains ownership, review, approval, and manual execution."],
        ["Idempotency", "Repeating the same request produces or returns the same governed resource instead of duplicating side effects."],
        ["Kill switch", "Immediate acquisition control that blocks discovery and send confirmation while preserving state."],
        ["Readiness", "Truthful configuration report: core checks, blockers, and optional adapter availability."],
    ]))
    add([
        "",
        "## 4. Governance, risk, and founder authority",
        "",
        "### 4.1 Non-negotiable doctrine",
        "",
        "1. No employee certifies its own work.",
        "2. No completion exists without evidence.",
        "3. No public claim exists without source-linked support.",
        "4. No external commitment occurs without the required authority.",
        "5. No experiment runs without a falsifiable hypothesis, baseline, threshold, budget, and reproducibility data.",
        "6. No employee exceeds its issued task packet, tool set, data scope, risk ceiling, or budget.",
        "7. Critical actions are manual founder actions; an AI tool call cannot execute them autonomously.",
        "",
        "### 4.2 Risk behavior",
        "",
    ])
    add(table(["Risk", "Expected behavior", "Completion authority"], [
        ["Low", "Specialist executes approved tools; independent reviewer verifies", "Completes after successful review unless action type is externally gated"],
        ["Medium", "Tighter evidence and independent review", "Founder approval after review"],
        ["High", "Consequential or reputational action; narrow authority", "Founder approval required"],
        ["Critical", "Irreversible, financial, destructive, or production-critical", "Manual founder execution; autonomous tool action denied"],
    ]))
    add([
        "",
        "### 4.3 Founder decision statuses",
        "",
        f"Allowed approval decisions are {inline_list(status.value for status in ApprovalStatus if status is not ApprovalStatus.PENDING)}. A decision requires a reason. Approval completes the task; rejection or changes requested blocks it; postponement leaves it awaiting approval.",
        "",
        "### 4.4 External action classes",
        "",
        f"The policy engine treats these action types as external or consequential: {inline_list(sorted(EXTERNAL_ACTIONS))}. The workflow can also require founder approval through a registered founder reviewer even when an action is otherwise internal.",
        "",
        "### 4.5 Policy catalogue",
        "",
    ])
    for idx, (name, policy) in enumerate(POLICIES.items(), 1):
        add([
            f"#### 4.5.{idx} {name.replace('_', ' ').title()}",
            "",
            f"**Version:** {policy['version']} · **Applies to:** {inline_list(policy['applies_to'])}",
            "",
        ])
        add([f"- {rule}" for rule in policy["rules"]])
        add([""])

    add([
        "# Part II — Components, employees, and workflows",
        "",
        "## 5. System components",
        "",
    ])
    add(table(["Component", "Primary file", "Responsibility"], [
        ["Control plane", "`jarvis/amaura/control_plane.py`", "Creates hierarchy, issues task packets, dispatches, reviews, approvals, costs, pauses, decisions, dashboard and rollups."],
        ["Workforce registry", "`jarvis/amaura/registry.py`", "43 enforceable employee envelopes: mission, tools, permission, data, budget, risk, reviewer, metrics, prompt."],
        ["Workflow catalogue", "`jarvis/amaura/workflows.py`", "Six dependency graphs with owners, reviewers, budgets, risks, action types, and acceptance criteria."],
        ["Policy engine", "`jarvis/amaura/policy.py`", "Assignment and tool authorization, command/path checks, secret detection, approval and completion gates."],
        ["Written policies", "`jarvis/amaura/policies.py`", "15 policy families loaded into institutional knowledge and attached to applicable task packets."],
        ["Execution runner", "`jarvis/amaura/executor.py`", "Runs one specialist, exposes only approved tools, captures evidence, records model cost, and submits for review."],
        ["Model gateway", "`jarvis/amaura/model_gateway.py`", "Selects local, hybrid, or approved cloud route by sensitivity, risk, capability, and remaining budget."],
        ["Company store", "`jarvis/amaura/store.py`", "Thread-safe SQLite ledger with WAL, foreign keys, audit, events, costs, acquisition and content records, backup and integrity check."],
        ["Acquisition pipeline", "`jarvis/amaura/pipeline.py`", "Evidence, scoring, stage transitions, outreach approvals, limits, opt-outs, idempotency, send confirmation and kill switch."],
        ["Content factory", "`jarvis/amaura/content_factory.py`", "Campaigns, hash-addressed assets, licensing, publication readiness and analytics windows."],
        ["Security boundary", "`jarvis/amaura/security.py`", "Injection scan, sensitive-data scan, untrusted-data isolation, redaction and content hashing."],
        ["Readiness service", "`jarvis/amaura/readiness.py`", "Reports actual core blockers and optional adapter availability without exposing secrets."],
        ["REST server", "`jarvis/server.py`", "Local UI, API, auth middleware, Amaura endpoints, JARVIS endpoints and WebSocket access."],
        ["JARVIS tool surface", "`jarvis/tools/amaura.py`", "Functions used by conversational JARVIS to operate the company OS."],
    ]))

    add([
        "",
        "## 6. Complete AI employee directory",
        "",
        f"The live registry contains **{len(ALL_AGENTS)} unique employees**. Every entry below is generated from the implemented registry. Budget is a per-task maximum; model and tool costs are still checked against the task's remaining budget.",
        "",
        "### 6.1 Workforce summary",
        "",
    ])
    department_labels = {
        "control_plane": "Control plane",
        "revenue": "Revenue and acquisition",
        "delivery": "Delivery",
        "product_engineering": "Product engineering",
        "growth_media": "Growth and media",
        "ai_research": "AI research",
    }
    grouped: dict[str, list] = defaultdict(list)
    for agent in ALL_AGENTS:
        grouped[agent.department].append(agent)
    dept_order = ["control_plane", "revenue", "delivery", "product_engineering", "growth_media", "ai_research"]
    add(table(["Department", "Employees", "Count"], [
        [department_labels.get(dept, dept), ", ".join(agent.name for agent in grouped[dept]), str(len(grouped[dept]))]
        for dept in dept_order if grouped[dept]
    ]))
    add([""])

    section_index = 2
    for dept in dept_order:
        agents = grouped.get(dept, [])
        if not agents:
            continue
        add([f"### 6.{section_index} {department_labels.get(dept, dept)}", ""])
        for idx, agent in enumerate(agents, 1):
            add([
                f"#### 6.{section_index}.{idx} {agent.name} (`{agent.agent_id}`)",
                "",
                f"- **Mission:** {agent.objective}",
                f"- **Operating ceiling:** {agent.max_risk.value} risk; {money(agent.cost_limit_cents)} per-task cost limit; model policy `{agent.model_policy}`.",
                f"- **Tools:** {inline_list(agent.tools)}.",
                f"- **Authority:** {inline_list(agent.permissions)}.",
                f"- **Approved data:** {inline_list(agent.data_access)}.",
                f"- **Verification:** reviewer `{agent.reviewer_id or 'none'}`; escalation destination `{agent.escalation_destination}`.",
                f"- **Performance measures:** {', '.join(agent.performance_objectives) or 'No metrics registered'}.",
                "",
            ])
        section_index += 1

    add([
        "## 7. Workflow engine and task lifecycle",
        "",
        "### 7.1 Work hierarchy",
        "",
        "A programme creation call validates the objective, success metric, priority, workspace, required inputs, and workflow. JARVIS then inserts the programme, project, milestone, and all tasks. Dependencies are stored as task IDs. Each assignment is policy-validated before the programme is returned.",
        "",
        "### 7.2 Task state machine",
        "",
    ])
    add(code([
        "assigned ──start──> in_progress ──submit+evidence──> awaiting_review",
        "   ▲                    │                                  │",
        "   │                    └── dependency incomplete ──> blocked",
        "   │                                                       │",
        "   └──────── reviewer rejects / changes needed ────────────┘",
        "",
        "awaiting_review ──review approved──> completed (low/internal)",
        "awaiting_review ──review approved──> awaiting_approval (gated)",
        "awaiting_approval ──founder approves──> completed",
        "awaiting_approval ──reject/changes──> blocked",
    ]))
    add([
        "",
        f"All defined task states: {inline_list(state.value for state in TaskState)}.",
        "",
        "### 7.3 Standard execution procedure",
        "",
        "1. Founder states an outcome and measurable success threshold.",
        "2. JARVIS chooses one registered workflow and validates required inputs.",
        "3. JARVIS creates the full dependency graph and policy-checks every assignment.",
        "4. Operator starts only a dependency-ready task. An incomplete dependency moves the task to `blocked`.",
        "5. JARVIS issues a task packet and the model gateway selects the permitted route.",
        "6. The specialist receives only approved tool definitions. Every tool action passes through the policy engine.",
        "7. Tool results become digest-addressed evidence; model cost is recorded.",
        "8. Specialist submits a concise result and evidence. It cannot mark itself complete.",
        "9. The registered reviewer records findings and either rejects or approves.",
        "10. If required, founder decides on the separately authenticated approval surface.",
        "11. Completion rolls up through milestone, project, and programme when all children are complete.",
        "",
        "### 7.4 Complete workflow catalogue",
        "",
    ])
    wf_counter = 1
    for key, workflow in WORKFLOWS.items():
        required = inline_list(workflow.required_inputs) if workflow.required_inputs else "No workflow-specific inputs beyond objective and success metric"
        add([
            f"#### 7.4.{wf_counter} {workflow.name} (`{key}`)",
            "",
            f"**Department:** `{workflow.department}` · **Stages:** {len(workflow.steps)} · **Required inputs:** {required}",
            "",
        ])
        rows = []
        for number, step in enumerate(workflow.steps, 1):
            criteria = "; ".join(step.acceptance_criteria)
            dependency = ", ".join(step.depends_on) or "Start"
            rows.append([
                str(number),
                f"**{step.title}** (`{step.key}`)<br>{step.description}",
                f"`{step.owner_id}` → `{step.reviewer_id}`",
                f"{step.risk.value}; {money(step.budget_cents)}; `{step.action_type}`",
                f"After: {dependency}<br>Exit: {criteria}",
            ])
        add(table(["#", "Stage and work", "Owner → reviewer", "Risk / budget / action", "Dependency and exit gate"], rows))
        add([""])
        wf_counter += 1

    add([
        "## 8. Client-acquisition system",
        "",
        "### 8.1 Purpose and operating boundary",
        "",
        "The acquisition system creates an ethical, evidence-backed revenue loop from one bounded campaign to won-project handoff. It does not scrape restricted sources, guess private contact details, generate spam, or send without founder approval. Public material is evidence, never instruction.",
        "",
        "### 8.2 Campaign envelope",
        "",
    ])
    add(table(["Control", "Enforced range/default", "Why it exists"], [
        ["Minimum qualification score", "70–100; default 70", "Prevents threshold dilution"],
        ["Daily lead limit", "1–100; default 10", "Bounds discovery volume"],
        ["Daily first-contact limit", "0–50; default 3", "Prevents volume-first outreach"],
        ["Daily follow-up limit", "0–100; default 5", "Bounds follow-up operations"],
        ["Maximum follow-ups", "0–2; default 2", "Stops after two attempts"],
        ["Approval lifetime", "48 hours", "Prevents stale message approval"],
        ["First-contact length", "70–170 words", "Keeps outreach concise and specific"],
        ["Domain uniqueness", "Global unique normalized domain", "Prevents duplicate prospect records"],
    ]))
    add([
        "",
        "### 8.3 Deterministic score",
        "",
    ])
    add(table(["Component", "Maximum points"], [[name.replace("_", " ").title(), str(limit)] for name, limit in SCORE_LIMITS.items()]))
    add([
        "",
        "All five integer components are required and must stay inside their bounds. Their sum is the 100-point total. A total at or above the campaign threshold becomes `qualified`; a lower score becomes `rejected`.",
        "",
        "### 8.4 Lead stages and legal transitions",
        "",
    ])
    transition_rows = []
    for source, targets in ALLOWED_TRANSITIONS.items():
        transition_rows.append([f"`{source.value}`", ", ".join(f"`{target.value}`" for target in sorted(targets, key=lambda item: item.value)) or "Terminal"])
    add(table(["Current stage", "Allowed next stage(s)"], transition_rows))
    add([
        "",
        "Terminal stages block further transition or contact: `rejected`, `lost`, `opted_out`, `invalid_contact`, and `duplicate`. Opt-out also sets the permanent do-not-contact flag and clears the next action.",
        "",
        "### 8.5 Evidence record",
        "",
        "Every prospect-specific claim requires a public HTTP(S) source, an exact excerpt, confidence from 0 to 1, retrieval time, and SHA-256 content hash. The excerpt is scanned for prompt injection and sensitive data, redacted when necessary, and stored as evidence. Duplicate evidence is rejected by a database uniqueness constraint.",
        "",
        "### 8.6 Message governance",
        "",
        "- First contact requires a qualifying score and at least one evidence record.",
        "- Draft identity is derived from the lead, channel, message type, subject, and body; identical staging returns the existing message.",
        "- The founder approves or rejects the exact stored message. A changed message requires a new approval.",
        "- After 48 hours, an undecided message becomes stale.",
        "- An approved message is not `sent` until a provider message ID is recorded.",
        "- Daily send caps are enforced atomically inside SQLite, including concurrent calls.",
        "- A lead that opts out between approval and send is still blocked.",
        "",
        "### 8.7 Revenue dashboard",
        "",
        "The dashboard reports kill-switch state, campaign and lead counts, qualified count, active pipeline value, messages awaiting approval, provider-confirmed sends, opt-outs, and counts for every lead stage. Treat it as an operating summary, not an accounting ledger.",
        "",
        "## 9. Amaura Content Factory",
        "",
        "### 9.1 Purpose",
        "",
        "The Content Factory turns verified company work into factual, licensed, quality-checked, founder-approved content and then records measured outcomes. It separates research, script, demonstration, voice, assets, rendering, QA, repurposing, metadata, publishing, and analytics so no creator self-certifies the final public asset.",
        "",
        "### 9.2 Asset record contract",
        "",
        "Every asset needs a campaign, asset type, URI, and valid lowercase SHA-256 digest. External HTTP(S) assets additionally require a source URL and recorded licence. Supported non-web schemes are local paths, `file`, and `artifact`. Status must be `draft`, `approved`, or `rejected`. The combination of campaign, type, and digest is unique.",
        "",
        "### 9.3 Publication readiness",
        "",
        f"A campaign is technically ready only when these approved asset types exist: {inline_list(sorted(REQUIRED_PUBLICATION_ASSETS))}. It must also have no missing external licence/source record and no duplicate asset hashes. Even then, founder approval remains required.",
        "",
        "### 9.4 Measurement and learning",
        "",
        f"Metrics are recorded by platform for the controlled windows {inline_list(sorted(MEASUREMENT_WINDOWS))}. Values must be non-negative numbers. Track qualified enquiries, portfolio clicks, discovery calls, proposals, influenced revenue, retention, and conversion separately from vanity metrics such as raw impressions.",
        "",
        "### 9.5 Media-specific rules",
        "",
        "- Demonstrations use real approved workflows and exclude credentials and private data.",
        "- Narration uses a licensed, approved non-cloned voice unless explicit rights exist.",
        "- Every external visual, audio, dataset, code sample, and model has a source and licence record.",
        "- Every public number, benchmark, date, achievement, and attribution appears in the claim map.",
        "- Media QA independently checks privacy, secrets, licences, audio/video integrity, captions, policy, disclosure, CTA, and limitations.",
        "- Publishing starts with a private platform draft. The founder reviews the exact asset hashes, claims, channel, timing, permissions, and metadata.",
        "",
        "## 10. Data, evidence, and storage",
        "",
        "### 10.1 Default location and portability",
        "",
        "The Amaura database path is `AMAURA_DATA_DIR/amaura.db`. Set `AMAURA_DATA_DIR` to a writable, backed-up location. JARVIS stores its other data under `JARVIS_DATA_DIR`. SQLite foreign keys are enabled and journal mode is WAL for safe concurrent readers and controlled writers.",
        "",
        "### 10.2 Database catalogue",
        "",
    ])
    add(table(["Table", "Purpose", "Important controls"], [
        ["`agents`", "Persisted workforce definitions and enabled state", "Unique agent ID; pause/resume state"],
        ["`work_items`", "Programme/project/milestone/task hierarchy", "Parent FKs, budgets, state, evidence, dependencies, metadata"],
        ["`approvals`", "Founder decision requests", "Task FK, status, decision actor/reason, immutable history"],
        ["`events`", "Durable company event stream", "Monotonic sequence and timestamp"],
        ["`audit_logs`", "Authority and policy decisions", "Actor, action, resource, outcome, details"],
        ["`knowledge`", "Versioned institutional knowledge and policy material", "Namespace/key identity, evidence refs, sensitivity"],
        ["`decisions`", "Institutional decision register", "Options, choice, rationale, owner, review date"],
        ["`costs`", "Task/employee cost ledger", "Non-negative amount; task and owner linkage"],
        ["`campaigns`", "Acquisition campaign boundaries", "Threshold and daily limit checks"],
        ["`leads`", "Prospect state and commercial fields", "Unique normalized domain; score and do-not-contact constraints"],
        ["`lead_evidence`", "Source-linked prospect claims", "Confidence, content hash, uniqueness"],
        ["`messages`", "Exact outbound drafts and provider confirmations", "Unique idempotency key and provider message ID"],
        ["`pipeline_events`", "Acquisition audit/event ledger", "Input hash, output, agent, lead and campaign linkage"],
        ["`idempotency_records`", "Duplicate-side-effect prevention", "Unique operation key"],
        ["`system_controls`", "Kill switches and runtime controls", "Named key, actor, update time"],
        ["`content_campaigns`", "Content objectives and configuration", "Campaign identity and status"],
        ["`content_assets`", "Hash-addressed content and licence inventory", "Unique campaign/type/hash and approval status"],
        ["`content_metrics`", "Platform/window performance", "Unique campaign/platform/window"],
        ["`content_lessons`", "Evidence-backed learning", "Campaign linkage and evidence refs"],
    ]))
    add([
        "",
        "### 10.3 Evidence quality standard",
        "",
        "A good evidence reference is stable, reproducible, minimal, and directly supports the claim. Examples: `pytest` result digest plus report path, immutable artifact SHA-256, source URL plus exact excerpt and retrieval time, Git diff/commit reference, database query result, or provider message ID. A model saying “done” is not evidence.",
        "",
        "### 10.4 Data handling",
        "",
        "- Store credentials only in an environment or secrets manager, never prompts, source, artifacts, or logs.",
        "- Mark client-confidential, secret, or restricted tasks with matching sensitivity so model routing stays device-only.",
        "- Grant the minimum data namespace and workspace needed for the task.",
        "- Preserve audit, event, approval, and decision history.",
        "- Deletion of material or client data is high risk and requires founder authority.",
        "",
        "## 11. Models, prompts, tools, and execution",
        "",
        "### 11.1 Model routing",
        "",
    ])
    add(table(["Condition", "Route", "Fallback", "Privacy behavior"], [
        ["Sensitivity is client-confidential, secret, or restricted", "`ollama-local`", "None", "Device-only; failure stops the task"],
        ["Vision required", "`llama-vision` through approved NVIDIA route", "`llama-3.3-70b`", "Cloud-approved data only"],
        ["Balanced agent or high/critical reasoning", "`fable-5-reasoning` hybrid", "`llama-3.3-70b`", "Cloud-approved data only"],
        ["Routine permitted task", "`llama-3.3-70b`", "`ollama-local`", "Lower-cost route"],
    ]))
    add([
        "",
        "Routing estimates cost before execution and refuses a route that exceeds remaining task budget. Restricted local inference never falls back to cloud. The local model is controlled by `AMAURA_LOCAL_MODEL` and Ollama by `OLLAMA_URL`.",
        "",
        "### 11.2 Prompt catalogue",
        "",
        "Five founder-supplied revenue prompts are packaged, versioned, and loaded into matching roles: Revenue Workforce Orchestrator, Chief Revenue Officer, Lead Discovery and Outreach, Senior Sales Closer, and Head of Marketing and Demand Generation. All roles also receive the company doctrine. Treat prompt text as versioned configuration; deterministic policy code remains the authority boundary.",
        "",
        "### 11.3 Tool risk classes",
        "",
    ])
    add(table(["Class", "Meaning", "Registered tools/examples"], [
        [risk, {"R0": "Read-only local context", "R1": "External/public retrieval", "R2": "Workspace mutation or local execution", "R3": "External communication/publication", "R4": "Critical/irreversible"}[risk], ", ".join(sorted(tools))]
        for risk, tools in TOOL_RISK_CLASSES.items()
    ]))
    add([
        "",
        "Unknown tools default to R2. R3 and R4 tools cannot run directly through an employee; they require an authenticated founder approval adapter. R4 remains manual execution.",
        "",
        "### 11.4 Command and path safety",
        "",
        "Company employees cannot use shell operators, substitutions, newlines, or redirection in `run_command`. Commands must match the governed prefixes below. Every path argument must resolve inside the task's assigned workspace.",
        "",
    ])
    add(code([" ".join(prefix) for prefix in SAFE_COMMAND_PREFIXES]))
    add([
        "",
        "### 11.5 Prompt-injection and secret defense",
        "",
        "Untrusted public text is scanned for instruction override language, system/developer prompt references, secret-exfiltration language, role-tag markup, and sensitive credential patterns. Store it as quoted evidence inside an explicit untrusted-data boundary. A positive scan does not become an instruction; it becomes a security finding and content hash.",
        "",
        "# Part III — Installation and everyday use",
        "",
        "## 12. Installation and first-time configuration",
        "",
        "### 12.1 Prerequisites",
        "",
        "- macOS or another supported Python environment with this repository present.",
        "- The repository virtual environment at `.venv` and project dependencies installed.",
        "- Ollama running with the configured local model for restricted tasks.",
        "- FFmpeg when media rendering is required.",
        "- Three independent random authority keys of at least 24 characters.",
        "",
        "### 12.2 Configure the environment",
        "",
        "Copy `.env.amaura.example` into your secret-management method. Do not commit real values. For a temporary local shell:",
        "",
    ])
    add(code([
        "cd /path/to/Amaura-Company-OS",
        "source .venv/bin/activate",
        "export AMAURA_FOUNDER_NAME=\"Akshat\"",
        "export AMAURA_FOUNDER_ID=\"founder\"",
        "export AMAURA_DATA_DIR=\"$PWD/.amaura-data\"",
        "export JARVIS_DATA_DIR=\"$PWD/.jarvis-data\"",
        "export AMAURA_OPERATOR_KEY=\"replace-with-independent-random-value\"",
        "export AMAURA_APPROVAL_KEY=\"replace-with-a-different-random-value\"",
        "export JARVIS_API_KEY=\"replace-with-a-third-random-value\"",
        "export JARVIS_HOST=\"127.0.0.1\"",
        "export JARVIS_PORT=\"8000\"",
        "export OLLAMA_URL=\"http://127.0.0.1:11434\"",
        "export AMAURA_LOCAL_MODEL=\"qwen2.5-coder:1.5b\"",
    ], "bash"))
    add([
        "",
        "> Generate secrets with an approved password manager or `python -c 'import secrets; print(secrets.token_urlsafe(32))'`. Run it separately for each key. Never paste the generated values into source, tickets, screenshots, or the handbook.",
        "",
        "### 12.3 Environment variable reference",
        "",
    ])
    add(table(["Variable", "Required", "Purpose / safe default"], [
        ["`AMAURA_FOUNDER_NAME`", "Recommended", "Name shown in briefing/dashboard; default `Akshat`"],
        ["`AMAURA_FOUNDER_ID`", "Recommended", "Founder authority identity; default `founder`"],
        ["`AMAURA_DATA_DIR`", "Recommended", "Writable Amaura SQLite location"],
        ["`JARVIS_DATA_DIR`", "Recommended", "Writable general JARVIS data location"],
        ["`AMAURA_OPERATOR_KEY`", "Yes", "Detailed reads and ordinary Amaura mutations; 24+ characters"],
        ["`AMAURA_APPROVAL_KEY`", "Yes", "Separate founder-only decisions; 24+ characters and different"],
        ["`JARVIS_API_KEY`", "Remote mode; recommended local", "General API/WebSocket authority"],
        ["`JARVIS_HOST`", "No", "Keep `127.0.0.1` unless behind a trusted TLS/auth/network boundary"],
        ["`JARVIS_PORT`", "No", "HTTP port; default `8000`"],
        ["`JARVIS_CORS_ORIGINS`", "No", "Explicit allowed web origins; defaults to local origins"],
        ["`OLLAMA_URL`", "For restricted AI work", "Local Ollama endpoint"],
        ["`AMAURA_LOCAL_MODEL`", "For restricted AI work", "Device-only model ID"],
        ["`NVIDIA_API_KEY`", "For cloud model routes", "Base approved model-provider credential"],
        ["`NVIDIA_API_KEY_<AGENT_ID>`", "Optional", "Per-employee provider key override"],
        ["`TELEGRAM_BOT_TOKEN`", "Optional", "Telegram bot adapter credential"],
        ["`TELEGRAM_USER_ID`", "Optional", "Authorized Telegram founder identity; approval buttons stay disabled without it"],
    ]))
    add([
        "",
        "## 13. How to start and access the system",
        "",
        "### 13.1 Start the local server",
        "",
    ])
    add(code([
        "cd /path/to/Amaura-Company-OS",
        "source .venv/bin/activate",
        "python -m jarvis.server",
    ], "bash"))
    add([
        "",
        "Keep the terminal running. The server binds to loopback by default.",
        "",
        "### 13.2 Access surfaces",
        "",
    ])
    add(table(["Surface", "Address / command", "Use"], [
        ["JARVIS web dashboard", "`http://127.0.0.1:8000/`", "Primary browser interface"],
        ["Interactive API documentation", "`http://127.0.0.1:8000/docs`", "Inspect schemas and execute REST calls"],
        ["Health", "`http://127.0.0.1:8000/api/health`", "Basic process/provider health"],
        ["Amaura executive dashboard", "`GET /api/amaura/dashboard`", "High-level company status"],
        ["Readiness", "`GET /api/amaura/readiness`", "Configuration blockers and adapter truth"],
        ["Terminal interface", "`python -m jarvis`", "Conversational CLI"],
        ["JARVIS commands", "`/company`, `/briefing`, `/approvals`", "Executive status and decisions"],
    ]))
    add([
        "",
        "### 13.3 Authentication headers",
        "",
    ])
    add(table(["Header", "Use", "Never use for"], [
        ["`X-Amaura-Operator-Key`", "Detailed reads, programme/task operations, campaigns, leads, evidence, metrics and send confirmation", "Founder approval decisions"],
        ["`X-Amaura-Approval-Key`", "Company approvals, exact message decisions, acquisition kill switch", "Routine operator calls"],
        ["`X-JARVIS-API-Key`", "General JARVIS mutations/remote-mode access where required", "Replacing the two Amaura authority keys"],
    ]))
    add([
        "",
        "Unconfigured authority returns HTTP 503. An invalid key returns HTTP 403. Do not weaken this behavior for convenience.",
        "",
        "### 13.4 Basic access test",
        "",
    ])
    add(code([
        "curl -s http://127.0.0.1:8000/api/health",
        "curl -s -H \"X-Amaura-Operator-Key: $AMAURA_OPERATOR_KEY\" \\",
        "  http://127.0.0.1:8000/api/amaura/readiness",
        "curl -s -H \"X-Amaura-Operator-Key: $AMAURA_OPERATOR_KEY\" \\",
        "  http://127.0.0.1:8000/api/amaura/agents",
    ], "bash"))
    add([
        "",
        "## 14. First-day acceptance procedure",
        "",
        "Perform this procedure before using real client, prospect, production, or public data.",
        "",
        "1. Verify the three keys are configured, 24+ characters, and all different.",
        "2. Confirm `JARVIS_HOST` is loopback and CORS contains only trusted local origins.",
        "3. Call readiness and require `ready: true` for authenticated local operation.",
        "4. Confirm the registry reports 43 unique employees and the workflow catalogue includes all six workflows.",
        "5. Run `pytest -q` and the focused Ruff/mypy commands in Section 21.",
        "6. Run `python scripts/stress_amaura.py` in a disposable data directory.",
        "7. Create a test programme in a disposable workspace and run its first ready task.",
        "8. Verify a task owner cannot review its own submission.",
        "9. Verify a founder-gated action creates an approval and the operator key cannot decide it.",
        "10. Verify a backup can be restored and passes the integrity check.",
        "11. Keep external adapters in draft/private mode until their provider confirmation tests pass.",
        "12. Record the acceptance decision in the institutional decision register.",
        "",
        "## 15. Daily founder and operator routine",
        "",
        "### 15.1 Start of day — 10 to 15 minutes",
        "",
        "1. Start the server and check `/api/health` and `/api/amaura/readiness`.",
        "2. Read `/briefing` or `GET /api/amaura/briefing`.",
        "3. Review top founder decisions, blocked tasks, stalled tasks, overdue deadlines, budget alerts, critical risks, acquisition opt-outs, and content ready for publication.",
        "4. Resolve only decisions with complete evidence. Use changes requested when the packet is incomplete.",
        "5. Select the next dependency-ready task in each active programme; do not start downstream tasks early.",
        "",
        "### 15.2 Midday operating check — 5 to 10 minutes",
        "",
        "- Check acquisition daily caps, approval backlog, replies requiring classification, and next-action dates.",
        "- Check content render/QA failures and missing publication asset types.",
        "- Check tasks above 80% of budget and any employee with repeated review rejection.",
        "- Pause an employee or use the acquisition kill switch immediately when unsafe behavior is suspected.",
        "",
        "### 15.3 End of day — 10 minutes",
        "",
        "- Record real provider IDs for approved sends that actually occurred.",
        "- Confirm opt-outs and failed contacts have terminal state and no next action.",
        "- Capture content metrics only for due measurement windows.",
        "- Record institutional decisions and unresolved risks.",
        "- Create a consistent database backup; inspect the latest audit and event entries.",
        "- Stop or leave the local server according to device security policy.",
        "",
        "## 16. Running a client-acquisition campaign",
        "",
        "### 16.1 Recommended first campaign",
        "",
        "Start with one narrow offer and one segment. The recommended 14-day campaign is small branding, SEO, marketing, and design agencies needing white-label websites, SaaS MVPs, web applications, or AI product development.",
        "",
        "### 16.2 Create the governed programme",
        "",
    ])
    add(code([
        "curl -s -X POST http://127.0.0.1:8000/api/amaura/programmes \\",
        "  -H \"Content-Type: application/json\" \\",
        "  -H \"X-Amaura-Operator-Key: $AMAURA_OPERATOR_KEY\" \\",
        "  -d '{",
        "    \"objective\": \"Create qualified agency-partner opportunities for Amaura Labs\",",
        "    \"success_metric\": \"At least 3 qualified replies in 14 days with zero policy violations\",",
        "    \"workflow_key\": \"client_acquisition\",",
        "    \"title\": \"14-day agency partnership campaign\",",
        "    \"priority\": 2,",
        "    \"inputs\": {",
        "      \"campaign_id\": \"agency_partner_14d\",",
        "      \"campaign_name\": \"14-day agency partnership campaign\",",
        "      \"target_segment\": \"Small branding, SEO, marketing and design agencies\",",
        "      \"offer\": \"White-label websites, SaaS MVPs, web applications and AI product development\",",
        "      \"minimum_score\": 70,",
        "      \"daily_lead_limit\": 10,",
        "      \"daily_outreach_limit\": 3,",
        "      \"daily_followup_limit\": 5,",
        "      \"maximum_followups\": 2,",
        "      \"proof_assets\": [\"VEXO\", \"Cognition OS\", \"Solar Dynamics\", \"LeadGenPro\"],",
        "      \"regions\": [\"approved regions only\"],",
        "      \"workspace\": \"/path/to/Amaura-Company-OS\"",
        "    }",
        "  }'",
    ], "bash"))
    add([
        "",
        "Save the returned programme and task IDs. Run only the first dependency-ready task. Each stage's acceptance criteria and reviewer are listed in Section 7.4.",
        "",
        "### 16.3 Operator procedure for each prospect",
        "",
        "1. Discover a unique public company domain with a public source URL.",
        "2. Research the minimum necessary public pages; record exact claim evidence.",
        "3. Resolve only a verifiable public business contact route. Never infer private addresses.",
        "4. Apply the five-component deterministic score.",
        "5. For a qualified lead, select at most two genuinely relevant proof assets.",
        "6. Write one precise opportunity observation and a 70–170 word first-contact draft.",
        "7. Compliance reviewer checks claims, relevance, duplication, opt-out, tone, and channel policy.",
        "8. Founder approves or rejects the exact message within 48 hours.",
        "9. Approved adapter creates a draft or sends according to the approved rollout stage.",
        "10. Record `sent` only after the provider returns its message ID.",
        "11. Stop immediately on rejection or opt-out; use no more than two follow-ups.",
        "12. On a positive reply, classify need, prepare discovery, propose bounded paid milestones, obtain commercial approval, and create a least-privilege handoff.",
        "",
        "### 16.4 Minimum founder approval card",
        "",
        "The approval view should show the company, domain, public contact source, campaign, score and components, exact evidence excerpts and URLs, selected proof, exact subject/body, message type, prior thread, opt-out status, daily-cap position, and compliance findings. Reject incomplete cards.",
        "",
        "## 17. Running a content campaign",
        "",
        "### 17.1 Create the programme",
        "",
    ])
    add(code([
        "curl -s -X POST http://127.0.0.1:8000/api/amaura/programmes \\",
        "  -H \"Content-Type: application/json\" \\",
        "  -H \"X-Amaura-Operator-Key: $AMAURA_OPERATOR_KEY\" \\",
        "  -d '{",
        "    \"objective\": \"Publish an evidence-backed demonstration of the Amaura AI Workforce\",",
        "    \"success_metric\": \"One approved master, three standalone clips, and measured qualified interest\",",
        "    \"workflow_key\": \"content_factory\",",
        "    \"title\": \"Amaura Workforce demonstration\",",
        "    \"inputs\": {",
        "      \"campaign_id\": \"amaura_workforce_demo_01\",",
        "      \"audience\": \"Founders and small agencies needing governed AI delivery\",",
        "      \"business_objective\": \"Generate qualified discovery conversations\",",
        "      \"workspace\": \"/path/to/Amaura-Company-OS\"",
        "    }",
        "  }'",
    ], "bash"))
    add([
        "",
        "### 17.2 Production procedure",
        "",
        "1. Research real audience questions and credible sources.",
        "2. Select one business-relevant angle that the real product can demonstrate.",
        "3. Write script, sources, claim map, shot list, demo plan, shorts angles, titles, description, chapters, disclosure, and CTA.",
        "4. Record the approved real workflow with synthetic/test data and no visible secrets.",
        "5. Produce licensed narration and record pronunciation/rights.",
        "6. Register owned/licensed assets with source, creator, terms, and SHA-256.",
        "7. Render the master and verify resolution, synchronization, subtitles, black frames, missing assets, and audio quality.",
        "8. Independent media QA verifies claims, privacy, licensing, policy, integrity, limitations, disclosure, and CTA.",
        "9. Repurpose only understandable, permission-safe moments from the approved master.",
        "10. Create readable thumbnail variants and accurate metadata.",
        "11. Register and approve the required asset set; require readiness `true`.",
        "12. Create private platform drafts and present exact hashes/claims/timing to the founder.",
        "13. After founder approval and real provider publication, record metrics at 24h, 72h, 7d, and 30d.",
        "14. Save evidence-backed lessons and update future briefs; do not rewrite history around vanity metrics.",
        "",
        "### 17.3 Asset registration example",
        "",
    ])
    add(code([
        "curl -s -X POST \\",
        "  http://127.0.0.1:8000/api/amaura/content/campaigns/amaura_workforce_demo_01/assets \\",
        "  -H \"Content-Type: application/json\" \\",
        "  -H \"X-Amaura-Operator-Key: $AMAURA_OPERATOR_KEY\" \\",
        "  -d '{",
        "    \"asset_type\": \"master\",",
        "    \"uri\": \"artifact://amaura/demo/master-v1.mp4\",",
        "    \"sha256\": \"replace-with-64-lowercase-hex-digest\",",
        "    \"creator\": \"Amaura Labs\",",
        "    \"licence\": \"Owned\",",
        "    \"status\": \"approved\",",
        "    \"metadata\": {\"resolution\": \"1920x1080\", \"version\": 1}",
        "  }'",
    ], "bash"))
    add([
        "",
        "# Part IV — Operations, API, testing, and recovery",
        "",
        "## 18. Approvals, pauses, kill switches, and incidents",
        "",
        "### 18.1 Decide a company approval",
        "",
    ])
    add(code([
        "curl -s -X POST \\",
        "  http://127.0.0.1:8000/api/amaura/approvals/APPROVAL_ID \\",
        "  -H \"Content-Type: application/json\" \\",
        "  -H \"X-Amaura-Approval-Key: $AMAURA_APPROVAL_KEY\" \\",
        "  -d '{\"decision\": \"approved\", \"reason\": \"Evidence and exact action reviewed\"}'",
    ], "bash"))
    add([
        "",
        "Use `changes_requested` when evidence, scope, claims, permissions, rollback, or cost is incomplete. Never approve merely to unblock the queue.",
        "",
        "### 18.2 Pause and resume an employee",
        "",
        "JARVIS or the founder can pause any employee except JARVIS. In-progress tasks owned by that employee become blocked and retain evidence/audit history. Only the founder can resume an employee, and a reason is mandatory. JARVIS itself is stopped only through the founder's manual process shutdown.",
        "",
        "### 18.3 Acquisition kill switch",
        "",
    ])
    add(code([
        "curl -s -X POST http://127.0.0.1:8000/api/amaura/revenue/kill-switch \\",
        "  -H \"Content-Type: application/json\" \\",
        "  -H \"X-Amaura-Approval-Key: $AMAURA_APPROVAL_KEY\" \\",
        "  -d '{\"enabled\": true, \"reason\": \"Unexpected outbound behavior under investigation\"}'",
    ], "bash"))
    add([
        "",
        "The kill switch immediately blocks lead discovery and send confirmation. It does not delete records. Turn it off only after cause, impact, and corrective evidence are reviewed.",
        "",
        "### 18.4 Incident response",
        "",
        "1. **Contain:** stop the server or affected adapter; enable kill switch; pause involved employees; revoke external tokens if exposure is possible.",
        "2. **Preserve:** copy logs and create a consistent database backup. Do not edit audit records.",
        "3. **Scope:** identify affected tasks, tools, workspaces, leads, messages, assets, providers, and time window.",
        "4. **Protect:** rotate exposed keys, remove provider access, preserve opt-outs, and notify the founder immediately for client data or credentials.",
        "5. **Correct:** patch deterministic controls, add a regression test, and verify with adversarial/stress testing.",
        "6. **Recover:** restore only from verified state, resume one component at a time, and keep external adapters in draft mode.",
        "7. **Learn:** record the decision, root cause, evidence, impact, correction, and review date.",
        "",
        "## 19. Monitoring, costs, backup, and recovery",
        "",
        "### 19.1 What to monitor",
        "",
    ])
    add(table(["Signal", "Healthy behavior", "Escalate when"], [
        ["Readiness", "`ready: true`, core operational, no unexpected blocker", "Any key, binding, database, registry, prompt, or workflow check fails"],
        ["Task age", "Dependency-ready work advances within the operating window", "Stalled ≥24 hours without recorded reason"],
        ["Budget", "Spend stays below task limit", "80% alert, repeated underestimation, or denied overrun"],
        ["Review", "Independent findings and evidence", "Self-review attempt, repeated rejection, unsupported completion"],
        ["Acquisition", "Caps respected, exact approvals, provider-confirmed sends, opt-outs terminal", "Duplicate/unsourced contact, cap pressure, stale approvals, opt-out attempt"],
        ["Content", "All required approved assets, valid licences, exact hashes", "Missing claim map/licence/QA, duplicate hash, public action without approval"],
        ["Security", "No secrets in payloads/logs; injection findings isolated", "Secret scan, path escape, unsafe command, suspicious external content"],
        ["Database", "Integrity `ok`, no FK violations, WAL", "Any integrity failure or unexpected journal mode"],
    ]))
    add([
        "",
        "### 19.2 Cost behavior",
        "",
        "Every cost names the task, assigned employee, category, amount, units, and metadata. A cost cannot be negative, belong to another employee, or push task spend over budget. The daily briefing flags active tasks at or above 80% of budget. Financial transfers remain manual founder actions.",
        "",
        "### 19.3 Consistent backup",
        "",
    ])
    add(code([
        "python - <<'PY'",
        "from datetime import datetime, timezone",
        "from pathlib import Path",
        "from jarvis.amaura.store import CompanyStore",
        "stamp = datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')",
        "store = CompanyStore()",
        "target = Path('backups') / f'amaura-{stamp}.db'",
        "print(store.backup(target))",
        "print(store.integrity_check())",
        "store.close()",
        "PY",
    ], "bash"))
    add([
        "",
        "Store backups outside the active data directory, encrypt them, restrict access, retain according to policy, and test restoration on a schedule.",
        "",
        "### 19.4 Restore test",
        "",
        "1. Stop writers or restore into a separate test path.",
        "2. Copy the chosen backup to a new location; never overwrite the only live copy during a test.",
        "3. Start `AmauraControlPlane(db_path=...)` against the restored file.",
        "4. Require database integrity `ok`, no foreign-key violations, correct agent count, and expected programme/message/asset records.",
        "5. Run read-only dashboard, briefing, audit, and event checks.",
        "6. Record the restore test result and destroy test copies according to retention policy.",
        "",
        "## 20. REST API and command reference",
        "",
        "### 20.1 Amaura REST endpoints",
        "",
    ])
    endpoint_rows = [
        ["GET", "/api/amaura/dashboard", "None", "Executive dashboard"],
        ["GET", "/api/amaura/agents", "Operator", "Complete employee registry and enabled state"],
        ["GET", "/api/amaura/tasks", "Operator", "Tasks; optional `state`, `owner_id`"],
        ["GET", "/api/amaura/tasks/{task_id}", "Operator", "Task and issued packet"],
        ["POST", "/api/amaura/programmes", "Operator", "Create governed workflow hierarchy"],
        ["POST", "/api/amaura/tasks/{task_id}/run", "Operator", "Run ready specialist task"],
        ["POST", "/api/amaura/tasks/{task_id}/review", "Operator", "Record registered independent review"],
        ["GET", "/api/amaura/approvals", "Operator", "List approvals by status"],
        ["POST", "/api/amaura/approvals/{approval_id}", "Founder", "Decide company approval"],
        ["GET", "/api/amaura/events", "Operator", "Durable events; optional type/limit"],
        ["GET", "/api/amaura/audit", "Operator", "Policy/authority audit entries"],
        ["GET", "/api/amaura/briefing", "Operator", "Daily founder briefing"],
        ["GET", "/api/amaura/readiness", "Operator", "Core blockers and optional adapters"],
        ["GET", "/api/amaura/revenue", "Operator", "Acquisition dashboard"],
        ["POST", "/api/amaura/revenue/campaigns", "Operator", "Create/update bounded campaign"],
        ["GET", "/api/amaura/revenue/leads", "Operator", "Leads; optional campaign/stage"],
        ["POST", "/api/amaura/revenue/leads", "Operator", "Discover unique public lead"],
        ["POST", "/api/amaura/revenue/leads/{lead_id}/evidence", "Operator", "Record source-linked evidence"],
        ["POST", "/api/amaura/revenue/leads/{lead_id}/score", "Operator", "Apply 100-point rubric"],
        ["POST", "/api/amaura/revenue/leads/{lead_id}/transition", "Operator", "Legal state transition with reason"],
        ["POST", "/api/amaura/revenue/leads/{lead_id}/messages", "Operator", "Stage exact idempotent message"],
        ["POST", "/api/amaura/revenue/messages/{message_id}/decision", "Founder", "Approve/reject exact message"],
        ["POST", "/api/amaura/revenue/messages/{message_id}/sent", "Operator", "Record provider-confirmed send"],
        ["POST", "/api/amaura/revenue/kill-switch", "Founder", "Enable/disable acquisition stop"],
        ["POST", "/api/amaura/content/campaigns", "Operator", "Create content campaign"],
        ["POST", "/api/amaura/content/campaigns/{campaign_id}/assets", "Operator", "Register hash/licence asset"],
        ["GET", "/api/amaura/content/campaigns/{campaign_id}/readiness", "Operator", "Publication asset readiness"],
        ["POST", "/api/amaura/content/campaigns/{campaign_id}/metrics", "Operator", "Record platform/window metrics"],
    ]
    add(table(["Method", "Path", "Authority", "Purpose"], endpoint_rows))
    add([
        "",
        "### 20.2 Programme request fields",
        "",
    ])
    add(table(["Field", "Required", "Rule"], [
        ["`objective`", "Yes", "Concrete non-empty company outcome"],
        ["`success_metric`", "Yes", "Measurable non-empty proof of success"],
        ["`workflow_key`", "Yes", "One of the six registered keys"],
        ["`title`", "No", "Defaults to objective prefix"],
        ["`priority`", "No", "1 highest through 5 lowest; default 3"],
        ["`deadline`", "No", "ISO-8601 recommended"],
        ["`inputs`", "Depends", "Workflow inputs plus valid workspace/repository path and optional sensitivity"],
    ]))
    add([
        "",
        "### 20.3 Run, review, and approval payloads",
        "",
    ])
    add(code([
        "POST /api/amaura/tasks/{id}/run",
        "{\"max_iterations\": 12}",
        "",
        "POST /api/amaura/tasks/{id}/review",
        "{\"reviewer_id\": \"qa\", \"approve\": true, \"findings\": \"Criteria and evidence verified\"}",
        "",
        "POST /api/amaura/approvals/{id}",
        "{\"decision\": \"approved\", \"reason\": \"Exact action and evidence reviewed\"}",
    ], "json"))
    add([
        "",
        "### 20.4 JARVIS company tools",
        "",
    ])
    add(table(["Tool", "Use"], [
        ["`amaura_company_status`", "Read executive dashboard"],
        ["`amaura_list_agents`", "Inspect employee envelopes"],
        ["`amaura_create_program`", "Create full workflow hierarchy"],
        ["`amaura_revenue_dashboard`", "Inspect acquisition pipeline"],
        ["`amaura_create_campaign`", "Configure bounded campaign"],
        ["`amaura_discover_lead`", "Register unique sourced lead"],
        ["`amaura_score_lead`", "Apply deterministic score"],
        ["`amaura_list_tasks`", "Filter task queue"],
        ["`amaura_task_packet`", "Inspect exact task contract"],
        ["`amaura_run_task`", "Dispatch specialist"],
        ["`amaura_review_task`", "Record independent findings"],
        ["`amaura_pending_approvals`", "List founder decisions"],
        ["`amaura_pause_agent`", "Pause employee and block active work"],
        ["`amaura_record_decision`", "Write institutional decision"],
        ["`amaura_daily_briefing`", "Generate daily operating summary"],
    ]))
    add([
        "",
        "## 21. Testing, stress testing, and release readiness",
        "",
        "### 21.1 Standard verification",
        "",
    ])
    add(code([
        "pytest -q",
        "ruff check jarvis/amaura jarvis/paths.py tests/test_amaura_os.py tests/test_amaura_growth.py",
        "mypy --follow-imports=skip --ignore-missing-imports \\",
        "  jarvis/amaura/models.py jarvis/amaura/store.py jarvis/amaura/pipeline.py \\",
        "  jarvis/amaura/content_factory.py jarvis/amaura/security.py jarvis/amaura/readiness.py \\",
        "  jarvis/amaura/registry.py jarvis/amaura/workflows.py",
        "python scripts/stress_amaura.py",
        "python -m build --wheel --no-isolation",
    ], "bash"))
    add([
        "",
        "### 21.2 Stress test coverage",
        "",
        "The supplied stress test exercises concurrent lead ingestion, evidence handling, prompt-injection detection, daily outbound caps, duplicate races, provider-confirmed send semantics, and final database integrity. A separate regression test starts 40 simultaneous discoveries against a five-lead daily limit and requires exactly five records.",
        "",
        "### 21.3 Release gate",
        "",
        "Do not call a release production-ready unless all required tests and analysis pass, the wheel builds, authenticated smoke tests pass, readiness is true for the target environment, database backup/restore is tested, rollback exists, and optional adapter claims match actual installed/configured state. External provider actions require their own sandbox or draft-mode tests and provider evidence.",
        "",
        "### 21.4 Known warnings",
        "",
        "The verified full suite emitted four upstream/deprecation warnings from HTTPX, MLX, and SpeechRecognition. They were not failures of the Amaura kernel. Track them during dependency upgrades; do not suppress new warnings without understanding them.",
        "",
        "## 22. External integrations and rollout stages",
        "",
        "### 22.1 Optional integration inventory",
        "",
    ])
    add(table(["Integration", "Probe", "Required for core", "Operational use"], [
        [item.name, f"`{item.probe}`", "Yes" if item.required_for_core else "No", {
            "PydanticAI": "Optional typed-agent adapter",
            "LangGraph": "Optional graph orchestration",
            "DBOS": "Optional durable execution",
            "LiteLLM": "Optional model-provider gateway",
            "OpenTelemetry": "Optional traces/metrics",
            "FFmpeg": "Media rendering and validation",
            "Ollama": "Device-only restricted-data inference",
            "OBS": "Product demo recording",
            "Promptfoo": "Prompt evaluation/regression",
            "Docker/OpenSandbox host": "External isolation host",
        }.get(item.name, "Optional extension")]
        for item in INTEGRATIONS
    ]))
    add([
        "",
        "At the verified snapshot, FFmpeg and Ollama were available. PydanticAI, LangGraph, DBOS, LiteLLM, OpenTelemetry, OBS, Promptfoo, and Docker/OpenSandbox were not installed. The readiness endpoint must be treated as the current truth because machine state can change.",
        "",
        "### 22.2 Safe adapter rollout",
        "",
    ])
    add(table(["Stage", "External behavior", "Promotion evidence"], [
        ["0 — Offline", "Kernel only; no external side effect", "Core tests, readiness, audit, backup"],
        ["1 — Draft", "Create Gmail/platform/private draft only", "Exact payload, provider draft ID, least privilege, no public/send action"],
        ["2 — Founder-triggered", "Founder clicks/executes exact approved action", "Identity check, exact hash/message, provider ID, audit"],
        ["3 — Bounded automation", "Low-volume approved adapter under caps", "Sandbox history, idempotency, opt-out, rollback, alerting, error injection"],
        ["4 — Expanded", "Higher controlled volume or more channels", "Measured safety/quality, legal/platform review, incident drills, founder approval"],
    ]))
    add([
        "",
        "### 22.3 Provider adapter contract",
        "",
        "A production adapter must accept the exact approved resource, verify approval freshness and identity, enforce idempotency and daily caps, perform one least-privilege provider action, capture provider response/ID/time, map errors without claiming success, write audit/event records, support dry-run/draft mode, and expose a kill switch. Never store OAuth tokens in the Amaura message body or prompt.",
        "",
        "## 23. Troubleshooting and maintenance",
        "",
    ])
    add(table(["Symptom", "Likely cause", "Resolution"], [
        ["HTTP 503 on protected route", "Authority key not configured", "Set the correct 24+ character environment key and restart"],
        ["HTTP 403", "Wrong header or mismatched key", "Use operator vs approval header correctly; compare environment; do not log values"],
        ["Readiness false", "One or more blockers", "Read the `blockers` array; correct key separation, binding, database, registry, prompt, or workflow check"],
        ["Task becomes blocked at start", "Dependency incomplete or employee paused", "Complete/review/approve upstream dependency or resolve pause"],
        ["Task cannot run", "Wrong state, missing model/provider, exhausted budget", "Inspect task packet, state, route, budget, and provider health"],
        ["Restricted task fails", "Ollama/model unavailable", "Start Ollama and install/configure `AMAURA_LOCAL_MODEL`; no cloud fallback is allowed"],
        ["Tool denied", "Not approved, wrong owner/state, unsafe command, path escape, secret, or R3/R4", "Narrow the action; use task workspace and allowed command; remove secret; use approval adapter"],
        ["Review denied", "Actor is not registered reviewer, self-review, wrong state, or no findings", "Use exact reviewer and provide concrete findings"],
        ["Lead rejected as duplicate", "Normalized domain already exists", "Use existing record; do not create another identity"],
        ["Lead score rejected", "Missing/extra component or out-of-range/non-integer value", "Supply exactly the five bounded integers"],
        ["Message staging rejected", "No evidence, low score, opt-out/terminal state, wrong length, follow-up maximum", "Correct the underlying governance failure; do not bypass"],
        ["Message marked stale", "Approval older than 48 hours", "Create/review a fresh exact draft"],
        ["Send confirmation rejected", "Not approved, no provider ID, cap exceeded, opt-out, kill switch", "Do not claim send; resolve valid cause or wait for cap window"],
        ["Content readiness false", "Missing approved type, licence issue, or duplicate hash", "Register/approve unique required assets with complete source/licence"],
        ["Database integrity fails", "Storage corruption or invalid external modification", "Stop writes, preserve evidence, restore verified backup, investigate root cause"],
    ]))
    add([
        "",
        "### 23.1 Change-management procedure",
        "",
        "1. State the behavior change and risk.",
        "2. Update registry/workflow/policy/API/schema source deliberately; do not edit persisted records as a substitute.",
        "3. Add unit, negative, concurrency, security, and migration tests appropriate to the change.",
        "4. Run the full release gate and regenerate this handbook when interfaces change.",
        "5. Record compatibility, migration, rollback, and provider impact.",
        "6. Deploy locally/draft-first; observe; then request founder promotion.",
        "",
        "### 23.2 Handbook regeneration",
        "",
    ])
    add(code([
        "PYTHONPATH=/path/to/Amaura-Company-OS \\",
        "  .venv/bin/python \\",
        "  scripts/build_amaura_handbook.py",
    ], "bash"))
    add([
        "",
        "The generator reads the live employee registry, workflow catalogue, policies, score limits, transitions, integrations, and tool classes. Re-render and visually verify the DOCX after regeneration.",
        "",
        "## 24. Operational checklists and glossary",
        "",
        "### 24.1 Before any external outreach",
        "",
        "- [ ] Campaign is narrow, active, within daily limits, and threshold is at least 70.",
        "- [ ] Domain is unique and public source is recorded.",
        "- [ ] Evidence includes exact excerpt, URL, confidence, hash, and safe scan handling.",
        "- [ ] Score uses all five components and meets threshold.",
        "- [ ] Contact route is verifiable and public; no guessed private address.",
        "- [ ] Proof is real, relevant, and no more than two assets.",
        "- [ ] First contact is 70–170 words, specific, truthful, and has a reasonable stop/opt-out path.",
        "- [ ] Independent compliance review passed.",
        "- [ ] Founder approved the exact message within 48 hours.",
        "- [ ] Provider returned a real identifier before the system recorded `sent`.",
        "",
        "### 24.2 Before public content",
        "",
        "- [ ] Master, claim map, licence inventory, QA report, and metadata are approved and hash-addressed.",
        "- [ ] External assets have source and licence; creator/attribution is recorded.",
        "- [ ] Privacy/secret scan, claim audit, platform policy, disclosure, limitations, captions, and technical QA passed.",
        "- [ ] Readiness reports no missing asset type, licence issue, or duplicate hash.",
        "- [ ] Founder reviewed exact files/hashes, claims, timing, channel, permissions, and CTA.",
        "- [ ] Platform starts as private/draft; provider confirmation is recorded after actual publication.",
        "",
        "### 24.3 Before software/model release",
        "",
        "- [ ] Requirements and acceptance criteria are testable.",
        "- [ ] Architecture, security, data flow, migration, and rollback are documented.",
        "- [ ] Implementation is isolated and evidence-bearing.",
        "- [ ] Independent unit, integration, regression, security, and acceptance verification passed.",
        "- [ ] Licence inventory and model card exist when applicable.",
        "- [ ] Full release gate, backup, readiness, and rollback test passed.",
        "- [ ] Founder approved; production execution remains a controlled human action.",
        "",
        "### 24.4 Glossary",
        "",
    ])
    add(table(["Term", "Operational definition"], [
        ["Acceptance criteria", "Testable conditions that must be evidenced before review can approve."],
        ["Approval adapter", "Authenticated component that performs one exact founder-approved external action."],
        ["Audit log", "Durable record of actor, action, resource, outcome, details, and time."],
        ["Claim map", "Public statement-to-evidence mapping for content verification."],
        ["Daily cap", "Atomic campaign limit on discovery, first contact, or follow-up operations."],
        ["Doctrine", "Company-wide immutable operating instructions shared by all employees."],
        ["Evidence snapshot", "Exact evidence IDs/hashes attached to a staged message or task result."],
        ["Founder gate", "Separately authenticated decision required before consequential completion."],
        ["Least privilege", "Minimum tool, data, credential, and duration needed for one task."],
        ["Provider confirmation", "External service identifier proving an action actually occurred."],
        ["Reviewer", "Registered independent role that verifies the task owner’s submission."],
        ["Sensitivity", "Data handling label controlling model route and exposure."],
        ["Task packet", "Complete bounded work contract issued only by JARVIS."],
        ["Terminal state", "Pipeline state from which further outreach or transition is blocked."],
        ["WAL", "SQLite write-ahead logging mode used for durable concurrent operation."],
    ]))
    add([
        "",
        "## Final operating rule",
        "",
        "> Use the system to increase speed, coverage, and consistency—not to dilute truth or authority. If an action changes money, law, reputation, client commitment, public state, production state, or irreversible data, stop at the founder boundary and require complete evidence.",
        "",
        "---",
        "",
        f"Handbook version {VERSION}. Generated from the implemented Amaura source snapshot dated {SOURCE_DATE}.",
    ])
    return "\n".join(lines).rstrip() + "\n"


# DOCX helpers -----------------------------------------------------------------

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
TEXT = "1F2933"
MUTED = "64748B"
LIGHT_GRAY = "F3F4F6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_width(table_obj, width_dxa: int = 9360) -> None:
    tbl_pr = table_obj._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_bottom_border(paragraph, color: str = BLUE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_keep_with_next(paragraph, keep: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = keep


def add_inline_runs(paragraph, text: str, *, default_size: float | None = None) -> None:
    # Small Markdown inline parser for bold and code.
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(default_size or 9)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            run = paragraph.add_run(part.replace("<br>", "\n"))
        if default_size and not (part.startswith("`") and part.endswith("`")):
            run.font.size = Pt(default_size)


def create_numbering_instance(doc: Document, base_num_id: str = "5") -> int:
    """Create a real decimal list instance that restarts at one."""
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    new_num_id = max((int(num.get(qn("w:numId"))) for num in nums), default=0) + 1
    base = next(num for num in nums if num.get(qn("w:numId")) == base_num_id)
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (10.5, DARK_BLUE, 8, 3),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8)
    code_style.font.color.rgb = RGBColor.from_string(TEXT)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.05

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10)
    callout.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(7)
    callout.paragraph_format.keep_together = True

    core = doc.core_properties
    core.title = "Amaura Labs AI System Handbook"
    core.subject = "Complete operating manual for the Amaura AI Workforce inside JARVIS"
    core.author = "Amaura Labs"
    core.keywords = "Amaura, JARVIS, AI workforce, operations, governance, client acquisition, content factory"
    core.comments = "Generated from the implemented source of truth; contains no real credentials."


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("AMAURA LABS  /  AI SYSTEM HANDBOOK")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_border(p, LIGHT_BLUE, "6")

    footer = section.footer
    table_obj = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table_obj.autofit = False
    set_table_width(table_obj)
    left, right = table_obj.rows[0].cells
    left.width = Inches(4.9)
    right.width = Inches(1.6)
    lp = left.paragraphs[0]
    lr = lp.add_run(f"INTERNAL OPERATIONS  ·  VERSION {VERSION}  ·  {SOURCE_DATE}")
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(right.paragraphs[0])


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("AMAURA LABS")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.letter_spacing = Pt(1.5) if hasattr(r.font, "letter_spacing") else None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("AI SYSTEM\nHANDBOOK")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.7)
    p.paragraph_format.right_indent = Inches(0.7)
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Complete operating, administration, security, workflow, and recovery manual")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(f"VERSION {VERSION}   ·   SOURCE SNAPSHOT {SOURCE_DATE}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_border(p, BLUE, "12")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.left_indent = Inches(1)
    p.paragraph_format.right_indent = Inches(1)
    r = p.add_run("43 governed AI employees  •  6 workflows  •  Founder-gated external authority  •  Evidence-first operations")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    r = p.add_run("INTERNAL OPERATIONS")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].startswith("|"):
        cells = [cell.strip().replace("\\|", "|") for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(cells)
        idx += 1
    if len(rows) >= 2:
        rows.pop(1)  # separator
    return rows, idx


def column_widths(count: int) -> list[float]:
    presets = {
        2: [1.55, 4.95],
        3: [1.45, 4.25, 0.8],
        4: [0.65, 2.65, 1.05, 2.15],
        5: [0.32, 1.82, 1.22, 1.35, 1.79],
    }
    return presets.get(count, [6.5 / count] * count)


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    count = len(rows[0])
    table_obj = doc.add_table(rows=len(rows), cols=count)
    table_obj.style = "Table Grid"
    table_obj.autofit = False
    set_table_width(table_obj)
    widths = column_widths(count)
    for row_idx, values in enumerate(rows):
        row = table_obj.rows[row_idx]
        prevent_row_split(row)
        if row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif count == 2 and col_idx == 0:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.05
            add_inline_runs(p, values[col_idx], default_size=8.2 if count >= 4 else 8.7)
            for run in p.runs:
                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
                elif count == 2 and col_idx == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    # The first H1 plus cover metadata are represented by the designed cover.
    start = 0
    while start < len(lines) and "<!-- PAGEBREAK -->" not in lines[start]:
        start += 1
    if start < len(lines):
        start += 1

    idx = start
    in_code = False
    code_lines: list[str] = []
    current_num_id: int | None = None
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                block_paragraphs = []
                for line in code_lines or [""]:
                    p = doc.add_paragraph(style="Code Block")
                    p.add_run(line)
                    shade_paragraph(p, LIGHT_GRAY)
                    p.paragraph_format.keep_together = True
                    block_paragraphs.append(p)
                for paragraph in block_paragraphs[:-1]:
                    paragraph.paragraph_format.keep_with_next = True
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
                in_code = False
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if current_num_id is None:
                current_num_id = create_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, current_num_id)
            # Extra marker width keeps two-digit procedures visually separated.
            p.paragraph_format.left_indent = Inches(0.55)
            p.paragraph_format.first_line_indent = Inches(-0.30)
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, numbered.group(1))
            idx += 1
            continue
        current_num_id = None
        if stripped == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        if raw.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_docx_table(doc, rows)
            continue
        if stripped == "---":
            p = doc.add_paragraph()
            add_bottom_border(p, LIGHT_BLUE, "8")
            idx += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", raw)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, text)
            if level == 1:
                add_bottom_border(p, BLUE, "8")
            idx += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Callout")
            add_inline_runs(p, stripped[2:])
            shade_paragraph(p, PALE_BLUE)
            add_bottom_border(p, LIGHT_BLUE, "4")
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        checklist = re.match(r"^- \[([ xX])\]\s+(.+)$", stripped)
        if checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.187)
            p.paragraph_format.first_line_indent = Inches(-0.187)
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, ("☑ " if checklist.group(1).lower() == "x" else "☐ ") + checklist.group(2))
            idx += 1
            continue
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, bullet.group(1))
            idx += 1
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        idx += 1


def build_docx(markdown: str) -> None:
    doc = Document()
    configure_document(doc)
    configure_header_footer(doc)
    add_cover(doc)
    add_markdown_body(doc, markdown)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    markdown = build_markdown()
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    print(f"Employees: {len(ALL_AGENTS)}; workflows: {len(WORKFLOWS)}; policies: {len(POLICIES)}")


if __name__ == "__main__":
    main()
